import os
import timeit
import tempfile
import tkinter as tk
import assemblyai as aai
from docx import Document
from tkinter import filedialog, messagebox
from moviepy.editor import VideoFileClip, AudioFileClip


aai.settings.api_key = 'd37006136c4b4ed6824f727ae9938020'
transcriber = aai.Transcriber()

def video_to_audio(video_file_path, audio_file_path):
    try:
        video_clip = VideoFileClip(video_file_path)
        audio_clip = video_clip.audio
        if audio_clip is not None:
            audio_clip.write_audiofile(audio_file_path, codec='pcm_s16le')
            video_clip.close()
            audio_clip.close()
        else:
            raise ValueError("The selected video file does not contain an audio track.")
    except Exception as e:
        raise e

def convert_audio_to_text(audio_path):
    # Convert audio file to a URL-compatible format
    if audio_path.endswith(".mp3"):
        audio_clip = AudioFileClip(audio_path)
        audio_path_wav = audio_path.replace(".mp3", ".wav")
        audio_clip.write_audiofile(audio_path_wav, codec='pcm_s16le')
        audio_clip.close()
        audio_path = audio_path_wav
    
    # Transcribe audio using AssemblyAI
    try:
        start_time = timeit.default_timer()
        transcript = transcriber.transcribe(audio_path)
        elapsed_time = timeit.default_timer() - start_time
        return transcript.text, elapsed_time
    except Exception as e:
        print(e)
        return str(e), None

def open_file_dialog(file_type):
    global video_path, audio_path
    if file_type == "video":
        video_path = filedialog.askopenfilename(title="Select a video file", filetypes=[("Video Files", "*.mp4 *.avi *.mov")])
    elif file_type == "audio":
        audio_path = filedialog.askopenfilename(title="Select an audio file", filetypes=[("Audio Files", "*.wav *.mp3")])

def convert_file():
    try:
        if option_var.get() == "Video to Text":
            if video_path:
                temp_audio_fd, temp_audio_path = tempfile.mkstemp(suffix=".wav")
                os.close(temp_audio_fd)

                video_to_audio(video_path, temp_audio_path)
                text, elapsed_time = convert_audio_to_text(temp_audio_path)

                text_box.delete(1.0, tk.END)
                text_box.insert(tk.END, text)
                if elapsed_time is not None:
                    time_label.config(text=f"Conversion Time: {elapsed_time:.2f} seconds")
                else:
                    time_label.config(text="Conversion Time: Error occurred")

                os.remove(temp_audio_path)
            else:
                messagebox.showwarning("Warning", "Please select a video file first!")
        elif option_var.get() == "Audio to Text":
            if audio_path:
                text, elapsed_time = convert_audio_to_text(audio_path)

                text_box.delete(1.0, tk.END)
                text_box.insert(tk.END, text)
                if elapsed_time is not None:
                    time_label.config(text=f"Conversion Time: {elapsed_time:.2f} seconds")
                else:
                    time_label.config(text="Conversion Time: Error occurred")
            else:
                messagebox.showwarning("Warning", "Please select an audio file first!")
    except Exception as e:
        messagebox.showerror("Error", str(e))

def save_to_word():
    text = text_box.get(1.0, tk.END).strip()
    if not text:
        messagebox.showwarning("Warning", "No text to save!")
        return
    
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if save_path:
        try:
            document = Document()
            document.add_heading("Converted Text", level=1)
            document.add_paragraph(text)
            document.save(save_path)
            messagebox.showinfo("Success", "Text saved to Word document successfully!")
        except Exception as e:
            messagebox.showerror("Error", str(e))

root = tk.Tk()
root.title("Video to Text Converter")
root.configure(bg="#f0f0f0")


# Configure grid layout
root.columnconfigure(0, weight=1)
root.columnconfigure(1, weight=1)
root.columnconfigure(2, weight=1)
root.rowconfigure(0, weight=0)
root.rowconfigure(1, weight=0)
root.rowconfigure(2, weight=1)
root.rowconfigure(3, weight=0)
root.rowconfigure(4, weight=0)

# Title label
title_label = tk.Label(root, text="Video to Text Converter", width=20,
         height=2, borderwidth=4, relief="ridge", font=("Helvetica", 20), bg="#003366", fg="white")
title_label.grid(row=0, column=0, columnspan=3, pady=10, sticky='n')

# Dropdown menu for conversion type
option_var = tk.StringVar(value="Video to Text")
option_menu = tk.OptionMenu(root, option_var, "Video to Text", "Audio to Text")
option_menu.grid(row=1, column=0, columnspan=3, pady=10, sticky='n')

# Text box for output
text_box = tk.Text(root, wrap=tk.WORD, width=100, height=20, bg="#ffffff", fg="#000000")
text_box.grid(row=2, column=0, columnspan=3, padx=5, pady=10, sticky='nsew')

# Time label
time_label = tk.Label(root, text="Conversion Time: N/A ", font=("Helvetica", 12), bg="#003366", fg="White")
time_label.grid(row=3, column=0, columnspan=3, pady=10, sticky='n')

# Button row
button_frame = tk.Frame(root, bg="#f0f0f0")
button_frame.grid(row=4, column=0, columnspan=3, pady=10, sticky='n')

# Open file button
open_button = tk.Button(button_frame, text="Open File", command=lambda: open_file_dialog("video" if option_var.get() == "Video to Text" else "audio"), bg="#4CAF50", fg="white", width=20, height=2)
open_button.pack(side=tk.LEFT, padx=5)

# Convert button
convert_button = tk.Button(button_frame, text="Convert", command=convert_file, bg="#008CBA", fg="white", width=20, height=2)
convert_button.pack(side=tk.LEFT, padx=5)

# Save to Word button
save_button = tk.Button(button_frame, text="Save to Word", command=save_to_word, bg="#FF5733", fg="white", width=20, height=2)
save_button.pack(side=tk.LEFT, padx=5)

# Exit button
exit_button = tk.Button(button_frame, text="Exit", command=root.destroy, bg="#FF0000", fg="white", width=20, height=2)
exit_button.pack(side=tk.LEFT, padx=5)

root.mainloop()
